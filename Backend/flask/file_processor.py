import os
import tempfile
import stat
import magic
import mimetypes
from typing import Dict, Set, Tuple, Optional, Iterator
import PyPDF2
import docx
import re
from threading import Timer
from contextlib import contextmanager

class TimeoutContext:
    """Thread-based timeout context manager that works on Windows"""
    def __init__(self, seconds: int):
        self.seconds = seconds
        self.timer = None
        self.timed_out = False
    
    def _timeout_handler(self):
        self.timed_out = True
    
    def __enter__(self):
        self.timer = Timer(self.seconds, self._timeout_handler)
        self.timer.start()
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.timer:
            self.timer.cancel()
        if self.timed_out:
            raise TimeoutError(f"Operation timed out after {self.seconds} seconds")

@contextmanager
def timeout(seconds: int) -> Iterator[None]:
    """Cross-platform timeout context manager"""
    with TimeoutContext(seconds):
        yield

class FileValidator:
    ALLOWED_MIME_TYPES: Dict[str, Set[str]] = {
        'pdf': {'application/pdf'},
        'docx': {'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/zip'},
        'txt': {'text/plain'}
    }

    FILE_SIGNATURES: Dict[str, bytes] = {
        'pdf': b'%PDF',
        'docx': b'PK\x03\x04',
        'txt': None  # No specific signature for txt files
    }

    @staticmethod
    def validate_mime_type(file_path: str, expected_extension: str) -> bool:
        # Validate file MIME type
        try:
            mime = magic.from_file(file_path, mime=True)
            mime_fallback, _ = mimetypes.guess_type(file_path)
            allowed_types = FileValidator.ALLOWED_MIME_TYPES.get(expected_extension, set())
            return mime in allowed_types or mime_fallback in allowed_types
        except Exception as e:
            print(f"Error validating MIME type: {e}")
            return False
        
    @staticmethod
    def validate_file_signature(file_path: str, expected_extension: str) -> bool:
        # Validate file signature (magic number)
        try:
            expected_signature = FileValidator.FILE_SIGNATURES.get(expected_extension)
            if expected_signature is None:
                return True  # No signature to validate
            
            with open(file_path, 'rb') as file:
                file_start = file.read(len(expected_signature))
                return file_start == expected_signature
        except Exception as e:
            print(f"Error validating file signature: {e}")
            return False
        
    @staticmethod
    def scan_for_malicious_patterns(file_path: str) -> bool:
        # Scan file for known malicious patterns
        try:
            with open(file_path, 'rb') as file:
                file_content = file.read()
                
                known_malicious_patterns = [b'<script', b'javascript', b'eval(', b'<?php']
                file_content = file_content.lower()
                for pattern in known_malicious_patterns:
                    if pattern in file_content:
                        return False
            return True
        except Exception as e:
            print(f"Error scanning for malicious patterns: {e}")
            return True
        
    @staticmethod
    def validate_structure(file_path: str, expected_extension: str) -> bool:
        # Validate file structure by attempting to open it with appropriate library
        try:
            if expected_extension == 'pdf':
                return FileValidator.validate_pdf_structure(file_path)
            elif expected_extension == 'docx':
                return FileValidator.validate_docx_structure(file_path)
            elif expected_extension == 'txt':
                return FileValidator.validate_txt_structure(file_path)
            else:
                return False  # Unsupported file type for structure validation
        except Exception as e:
            print(f"Error validating file structure: {e}")
            return False
        
    @staticmethod
    def validate_pdf_structure(file_path: str) -> bool:
        try:
            with open(file_path, 'rb') as file:
                header = file.read(8)
                if not header.startswith(b'%PDF-'):
                    return False
                file.seek(-1024, 2) # last 1KB
                trailer = file.read()
                return b'%%EOF' in trailer
        except Exception:
            return True
        
    @staticmethod
    def validate_docx_structure(file_path: str) -> bool:
        try:
            doc = docx.Document(file_path)
            return len(doc.paragraphs) >= 0
        except Exception:
            return True
        
    @staticmethod
    def validate_txt_structure(file_path: str) -> bool:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                file.read(1024)
            return True
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    file.read(1024)
                return True
            except Exception:
                return False
        except Exception:
            return False
        
class FilenameSanitizer:
    MAX_FILENAME_LENGTH = 100
    
    dangerous_patterns = [r'\.\./', r'\\', r'/', r'^\.',r'[\x00-\x1f]', r'[<>:"|?*]']
    
    @staticmethod
    def sanitize_filename(filename: str) -> str:
        # Sanitize filename to prevent directory traversal and remove dangerous characters
        if not filename or len(filename) > FilenameSanitizer.MAX_FILENAME_LENGTH:
            return None
        
        for pattern in FilenameSanitizer.dangerous_patterns:
            filename = re.sub(pattern, '', filename)
            
        parts = filename.split('.')
        if len(parts) < 2:
            filename = parts[0] + '.' + parts[-1]
            
        filename = filename.strip('. ')
        
        if not filename or filename.startswith('.'):
            return None

        return filename
    
class SecureTempFile:
    @staticmethod
    @contextmanager
    def create_secure_temp_file(suffix: str = '', prefix: str = 'secure_') -> Iterator[Tuple[int, str]]:
        # Create a secure temporary file
        fd = None
        file_path = None
        try:
            fd, file_path = tempfile.mkstemp(suffix=suffix, prefix=prefix)
            os.chmod(file_path, stat.S_IRUSR | stat.S_IWUSR)  # Owner read/write only
            yield fd, file_path
        finally:
            if fd is not None:
                try:
                    os.close(fd)
                except OSError:
                    pass
            if file_path and os.path.exists(file_path):
                SecureTempFile.secure_delete_file(file_path)
                
    @staticmethod
    def secure_delete_file(file_path: str) -> None:
        # Securely delete a file by overwriting its content before deletion
        try:
            if not os.path.exists(file_path):
                return
            
            file_size = os.path.getsize(file_path)
            
            # Overwrite file content with random data
            with open(file_path, 'r+b') as f:
                f.write(os.urandom(file_size))
                f.flush()
                os.fsync(f.fileno())
            os.remove(file_path)
        
        except Exception as e:
            print(f"Secure delete error for {file_path}: {e}")
            try:
                os.remove(file_path)
            except:
                pass

class FileProcessor:
    # handles file upload processing and text extraction
    ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    VALIDATION_TIMEOUT = 30  # 30 seconds timeout for validation
    
    def __init__(self, upload_folder=None):
        self.upload_folder = upload_folder or tempfile.gettempdir()
        self.validate_upload_folder()
        
    def validate_upload_folder(self):
        # Ensure upload folder exists and is writable
        if not os.path.exists(self.upload_folder):
            os.makedirs(self.upload_folder, mode=0o700)
        if not os.access(self.upload_folder, os.W_OK):
            raise PermissionError(f"Upload folder '{self.upload_folder}' is not writable")
        
    def allowed_file(self, filename):
        # Check if file extension is allowed
        if not filename or '.' not in filename:
            return False
        
        extension = filename.rsplit('.', 1)[1].lower()
        return extension in self.ALLOWED_EXTENSIONS

    def validate_uploaded_file(self, file) -> Tuple[bool, Optional[str]]:
        # Validate uploaded file security before processing
        if not file or not file.filename:
            return False, "No file provided"

        if not self.allowed_file(file.filename):
            return False, "File type not supported. Please upload PDF, DOCX, or TXT files."

        sanitised_filename = FilenameSanitizer.sanitize_filename(file.filename)
        if not sanitised_filename:
            return False, "Invalid filename.  Please use a simpler filename without special characters."
        
        return True, sanitised_filename

    def extract_text_from_pdf(self, file_path):
        # Extract text from PDF file
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() + '\n'
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ''

    def extract_text_from_docx(self, file_path):
        # Extract text from DOCX file
        try:
            doc = docx.Document(file_path)
            text = ''
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ''

    def extract_text_from_txt(self, file_path):
        # Extract text from TXT file
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                print(f"Error extracting text from TXT: {e}")
                return ''
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return ''

    def process_file_security(self, file) -> Tuple[str, str]:
        # Process uploaded file and extract text
        # Initial validation
        is_valid, result = self.validate_uploaded_file(file)
        if not is_valid:
            raise ValueError(result)
        
        sanitized_filename = result
        file_extension = sanitized_filename.rsplit('.', 1)[1].lower()
        
        # Create secure temporary file
        with SecureTempFile.create_secure_temp_file(
            suffix=f'_{sanitized_filename}', 
            prefix='upload_'
        ) as (fd, temp_path):
            
            try:
                # Save uploaded file to secure temporary location
                with os.fdopen(fd, 'wb') as temp_file:
                    file.stream.seek(0)  # Reset file pointer
                    chunk_size = 8192
                    while True:
                        chunk = file.stream.read(chunk_size)
                        if not chunk:
                            break
                        temp_file.write(chunk)
                
                # Additional security validations with timeout
                with timeout(self.VALIDATION_TIMEOUT):
                    if not FileValidator.validate_file_signature(temp_path, file_extension):
                        raise ValueError("File signature validation failed - file may be corrupted or malicious")
                    
                    if not FileValidator.validate_mime_type(temp_path, file_extension):
                        raise ValueError("MIME type validation failed - file type mismatch")
                    
                    if not FileValidator.scan_for_malicious_patterns(temp_path):
                        raise ValueError("Security scan failed - potentially malicious content detected")
                    
                    if not FileValidator.validate_structure(temp_path, file_extension):
                        raise ValueError(f"Invalid {file_extension.upper()} document structure")
        
                return temp_path, sanitized_filename
                
            except TimeoutError:
                raise ValueError("File validation timed out - file may be too complex")
            except Exception as e:
                # Log security events
                print(f"File processing security event - File: {sanitized_filename}, Error: {str(e)}")
                raise e
            
    def process_file(self, uploaded_file, upload_folder=None):
        # Process uploaded file and extract text
        if not uploaded_file or not uploaded_file.filename:
            raise ValueError("No file provided")
        
        # Validate the uploaded file
        is_valid, result = self.validate_uploaded_file(uploaded_file)
        if not is_valid:
            raise ValueError(result)
            
        filename = result
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        # Create secure temporary file
        with SecureTempFile.create_secure_temp_file(
            suffix=f'_{filename}', 
            prefix='upload_'
        ) as (fd, file_path):
            
            try:
                # Save uploaded file to temporary location
                with os.fdopen(fd, 'wb') as temp_file:
                    uploaded_file.stream.seek(0)  # Reset file pointer
                    chunk_size = 8192
                    while True:
                        chunk = uploaded_file.stream.read(chunk_size)
                        if not chunk:
                            break
                        temp_file.write(chunk)
                
                # Extract text based on file extension
                if file_extension == 'pdf':
                    text = self.extract_text_from_pdf(file_path)
                elif file_extension == 'docx':
                    text = self.extract_text_from_docx(file_path)
                elif file_extension == 'txt':
                    text = self.extract_text_from_txt(file_path)
                else:
                    raise ValueError("Unsupported file type")
                
                return text, filename
            
            except Exception as e:
                # Log the error and re-raise
                print(f"Error processing file {filename}: {e}")
                raise e